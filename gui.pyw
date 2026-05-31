import win32com.client
import os
import tkinter as tk
from tkinter import ttk, filedialog, scrolledtext, messagebox
import threading
import shutil
import fitz  # PyMuPDF
import concurrent.futures
import multiprocessing
import math
from datetime import timedelta
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
import subprocess
import sys
from tkinterdnd2 import TkinterDnD, DND_FILES 
import ctypes
import time
import queue
import gc
#像素太长的图片可能会被库拒绝处理，可自行设置 PIL.Image.MAX_IMAGE_PIXELS = None
# --- Worker helpers for parallel processing (must be top-level for multiprocessing on Windows) ---
def _normalize_pixel(p):
    if isinstance(p, int):
        return (p, p, p)
    try:
        return tuple(p)[:3]
    except Exception:
        return (0, 0, 0)

def _find_segments_from_image(image, config):
    """Return segments as list of {'start': int, 'end': int} from a PIL Image using config dict."""
    width, height = image.size
    pixels = list(image.getdata())

    blank_height = int(config.get('blank_height', 10))
    threshold = int(config.get('threshold', 200))
    lr_enabled = bool(config.get('lr_threshold_enabled', False))
    lr_percent = int(config.get('lr_threshold_percent', 8))

    segments = []
    start = None
    blanks = 0

    for y in range(height):
        # build row based on lr settings (only affects blank detection)
        if lr_enabled and lr_percent > 0:
            left_cut = int(width * (lr_percent / 100.0))
            right_cut = width - left_cut
            if left_cut >= right_cut:
                row = [pixels[y * width + x] for x in range(width)]
            else:
                row = [pixels[y * width + x] for x in range(left_cut, right_cut)]
        else:
            row = [pixels[y * width + x] for x in range(width)]

        # normalize pixels and check blank condition
        is_blank = True
        for pix in row:
            r, g, b = _normalize_pixel(pix)
            if (r + g + b) / 3 < threshold:
                is_blank = False
                break

        if is_blank:
            blanks += 1
            if start is not None and blanks > blank_height:
                segments.append({'start': start, 'end': y - blank_height})
                start = None
        else:
            if start is None:
                start = max(0, y - 2)
            blanks = 0

    if start is not None:
        segments.append({'start': start, 'end': height})

    return segments

def _worker_segment_images(items, config, output_dir, dpi, progress_proxy=None, process_id=0, log_queue=None):
    """Worker function to segment existing full-page images.
    `items` is a list of (page_num, image_path).
    Returns list of (page_num, [segment_paths])."""
    from PIL import Image
    import os

    results = []
    os.makedirs(output_dir, exist_ok=True)

    for page_num, img_path in items:
        seg_paths = []
        try:
            msg = f"[segment-worker {process_id}] processing page {page_num}: {img_path}"
            print(msg)
            if log_queue is not None:
                try:
                    log_queue.put(('info', msg))
                except Exception:
                    pass
            with Image.open(img_path) as img:
                width, height = img.size
                segments = _find_segments_from_image(img, config)
                counter = 1
                for seg in segments:
                    seg_height = seg['end'] - seg['start']
                    if seg_height >= int(config.get('min_height', 10)):
                        segment_img = img.crop((0, seg['start'], width, seg['end']))
                        output_path = os.path.join(output_dir, f"segment_{page_num}_{counter}.png")
                        segment_img.save(output_path, dpi=(dpi, dpi))
                        seg_paths.append(output_path)
                        counter += 1
            # hint GC inside worker to free image memory earlier
            try:
                gc.collect()
            except Exception:
                pass
        except Exception:
            seg_paths = []
            try:
                if log_queue is not None:
                    log_queue.put(('error', f"[segment-worker {process_id}] error processing page {page_num}"))
            except Exception:
                pass
        msg2 = f"[segment-worker {process_id}] finished page {page_num}, segments: {len(seg_paths)}"
        print(msg2)
        if log_queue is not None:
            try:
                log_queue.put(('info', msg2))
            except Exception:
                pass
        results.append((page_num, seg_paths))

        # notify progress (increment by 1 page processed)
        try:
            if progress_proxy is not None:
                progress_proxy.value = int(progress_proxy.value) + 1
        except Exception:
            try:
                progress_proxy.put(1)
            except Exception:
                pass

    return results

def _render_page_to_image(pdf_path, page_idx, output_path, dpi, log_queue=None, process_id=0):
    """Render a single PDF page to an image file. Runs in worker process."""
    try:
        import fitz
        import os
        import gc

        doc = fitz.open(pdf_path)
        page = doc.load_page(page_idx)
        pix = page.get_pixmap(dpi=dpi)
        pix.save(output_path)
        # cleanup
        try:
            del pix
            del page
            doc.close()
        except Exception:
            pass
        gc.collect()

        msg = f"[render-worker] Rendered page {page_idx+1} -> {output_path} (pid:{os.getpid()})"
        print(msg)
        if log_queue is not None:
            try:
                log_queue.put(('info', msg))
            except Exception:
                pass

        return (page_idx + 1, output_path)
    except Exception as e:
        try:
            if log_queue is not None:
                log_queue.put(('error', f"Failed rendering page {page_idx+1}: {e}"))
        except Exception:
            pass
        return (page_idx + 1, None)

class PDFProcessorApp:
    def __init__(self, master):
        self.master = master
        master.title("FenScribe ~ PDF空白删")
        master.geometry("800x688")

        # 初始化配置参数
        self.config = {
            'threshold': 200,
            'blank_height': 10,
            'min_height': 10,
            'bottom_margin': 0,
            'dpi': 300,
            'lr_threshold_enabled': False,
            'lr_threshold_percent': 8,
            'template_docx': "",
            'save_interval': 688
        }

        # 创建界面组件
        self.create_widgets()
        self.load_templates()
        self.load_macros()

        # 状态变量
        self.is_processing = False
        self.progress = 0
        self._last_progress_update = 0
        self._pending_progress = None
        self._progress_flush_scheduled = False
        self._progress_throttle_seconds = 5

        # 绑定拖放事件
        self.master.drop_target_register(DND_FILES)  
        self.master.dnd_bind('<<Drop>>', self.handle_drop)  

    def handle_drop(self, event):
        """处理文件拖放事件"""
        files = event.data.strip('{}').split('} {')
        if len(files) > 1:
            messagebox.showerror("Error", "Only one PDF file can be dragged and dropped at a time")
            return
        file_path = files[0].strip()
        if os.path.isdir(file_path) or file_path.lower().endswith(('.pdf', '.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp', '.gif')):
            self.pdf_path.set(file_path)
            self.log(f"Selected file/folder: {file_path}")
        else:
            messagebox.showerror("Error", "Please drop a PDF, image, or folder containing images")

    def create_widgets(self):
        # 主容器
        main_frame = ttk.Frame(self.master)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)

        # 文件选择区域
        file_frame = ttk.LabelFrame(main_frame, text="Step1: Select the PDF file")
        file_frame.pack(fill="x", pady=5)

        self.pdf_path = tk.StringVar()
        ttk.Entry(file_frame, textvariable=self.pdf_path, width=60).pack(side="left", padx=5, fill="x", expand=True)
        btn_frame = ttk.Frame(file_frame)
        btn_frame.pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Browse...", command=self.select_pdf).pack(side="left")
        ttk.Button(btn_frame, text="Browse Folder...", command=self.select_folder).pack(side="left", padx=(5,0))

        # 参数设置区域
        param_frame = ttk.LabelFrame(main_frame, text="Step2. Set the processing parameters")
        param_frame.pack(fill="x", pady=5)

        ttk.Label(param_frame, text="Blank threshold (20-255):").grid(row=0, column=0, padx=5, sticky="w")
        self.threshold = ttk.Spinbox(param_frame, from_=20, to=255, width=8)
        self.threshold.set(200)
        self.threshold.grid(row=0, column=1, padx=5)

        ttk.Label(param_frame, text="DPI (72-600):").grid(row=0, column=2, padx=5, sticky="w")
        self.dpi = ttk.Spinbox(param_frame, from_=72, to=600, width=8)
        self.dpi.set(300)
        self.dpi.grid(row=0, column=3, padx=5)

        ttk.Label(param_frame, text="min_height:").grid(row=0, column=4, padx=5, sticky="w")
        self.min_height = ttk.Spinbox(param_frame, from_=1, to=300, width=8)
        self.min_height.set(10)
        self.min_height.grid(row=0, column=5, padx=5)

        ttk.Label(param_frame, text="blank_height:").grid(row=0, column=6, padx=5, sticky="w")
        self.blank_height = ttk.Spinbox(param_frame, from_=1, to=300, width=8)
        self.blank_height.set(10)
        self.blank_height.grid(row=0, column=7, padx=5)

        # Advanced (hidden) options toggle
        self.advanced_visible = False
        self.lr_enabled_var = tk.BooleanVar(value=False)
        self.lr_percent_var = tk.IntVar(value=8)
        self.enable_log_var = tk.BooleanVar(value=True)
        self.parallel_var = tk.BooleanVar(value=True)

        self.advanced_btn = ttk.Button(param_frame, text="Advanced ▶", command=self.toggle_advanced)
        self.advanced_btn.grid(row=1, column=0, padx=5, pady=5, sticky='w')

        self.advanced_frame = ttk.Frame(param_frame)
        ttk.Checkbutton(self.advanced_frame, text="启用 左右阈值", variable=self.lr_enabled_var).grid(row=0, column=0, sticky='w', padx=5)
        ttk.Label(self.advanced_frame, text="左右各 (%)").grid(row=0, column=1, sticky='w')
        self.lr_spinbox = ttk.Spinbox(self.advanced_frame, from_=0, to=49, width=8, textvariable=self.lr_percent_var)
        self.lr_spinbox.grid(row=0, column=2, padx=5)
        ttk.Checkbutton(self.advanced_frame, text="Enable parallel processing (faster)", variable=self.parallel_var).grid(row=0, column=3, sticky='w', padx=10)
        ttk.Label(self.advanced_frame, text="注: 仅影响空白行识别，不影响裁切输出").grid(row=1, column=0, columnspan=3, sticky='w', padx=5)
        self.no_insert_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(self.advanced_frame, text="不插入 Word（仅保存图片）", variable=self.no_insert_var).grid(row=2, column=0, columnspan=3, sticky='w', padx=5, pady=(4,0))
        ttk.Checkbutton(self.advanced_frame, text="显示日志 (高级)", variable=self.enable_log_var, command=self._toggle_log_frame).grid(row=3, column=0, columnspan=3, sticky='w', padx=5, pady=(4,0))

        # 模板选择区域
        template_frame = ttk.LabelFrame(main_frame, text="Step3. Select an output template")
        template_frame.pack(fill="x", pady=5)

        self.template_var = tk.StringVar()
        self.template_combobox = ttk.Combobox(template_frame, textvariable=self.template_var, state="readonly")
        self.template_combobox.pack(pady=5, fill="x", padx=5)

        macro_frame = ttk.LabelFrame(main_frame, text="Step4. Select a macro (optional)")
        macro_frame.pack(fill="x", pady=5)

        self.macro_var = tk.StringVar()
        self.macro_combobox = ttk.Combobox(macro_frame, textvariable=self.macro_var, state="readonly")
        self.macro_combobox.pack(pady=5, fill="x", padx=5)

        # 进度条
        self.progress_bar = ttk.Progressbar(main_frame, orient="horizontal", mode="determinate")
        self.progress_bar.pack(fill="x", pady=5)
        self.progress_label = ttk.Label(main_frame, text="")
        self.progress_label.pack(fill="x")

        # 日志区域
        self.log_frame = ttk.LabelFrame(main_frame, text="Process logs")
        self.log_area = scrolledtext.ScrolledText(self.log_frame, wrap=tk.WORD, height=6)
        self.log_area.pack(fill="both", expand=True)

        # 控制按钮
        self.control_frame = ttk.Frame(main_frame)
        self.control_frame.pack(pady=10)

        self.process_btn = ttk.Button(self.control_frame, text="Start processing", command=self.start_processing)
        self.process_btn.pack(side="left", padx=5)

        ttk.Button(self.control_frame, text="Open Output Folder", command=self.openf).pack(side="left", padx=5)
        ttk.Button(self.control_frame, text="DoubleColumnCut", command=self.open_double_column).pack(side="left", padx=5)
        ttk.Button(self.control_frame, text="Clear the log", command=self.clear_log).pack(side="left", padx=5)
        ttk.Button(self.control_frame, text="Quit", command=self.master.quit).pack(side="right", padx=5)

        welcome_text = """
Welcome to FenScribe!!    https://ordylan.com/FenScribe
欢迎: 以下是参数说明: 
1. threshold: 判断某一行是否为空白的亮度阈值 [即越大被误删概率越低]
   将像素RGB值转为灰度(取平均值), 当整行像素的平均灰度 ≥ 该值时视为空白行 

2. dpi: PDF转图片的分辨率  |   PS: 不采用跨行排版原因: 技术有限 格式错乱概率更大. 

3. min_height: 有效内容过滤阈值
   切割后的段落高度 ≥ 该值才会被保留, 否则视为无效内容

4. blank_height: 识别段落分隔的基准线
   当连续空白行数 ≥ 该值时, 认为前后内容属于不同段落
注意: 请先手动裁边pdf哦哦! pdf扫描版需要矫正角度; 多栏pdf请分成单栏进行切割(下有不佳的工具)哦! 
这是第四代代码 已经支持多线程加速!  
注: 宏功能(useless)会关闭已打开的word 注意保存工作! (暂未修复)
计划--添加传参自动处理功能  隐藏gui 方便其他程序调用

2026年5月31日  还要一周高考   高考必胜!! 

"""
        if self.enable_log_var.get():
            self.log(welcome_text)
            try:
                self._toggle_log_frame()
            except Exception:
                pass

    def open_double_column(self):
        py_path = os.path.join(os.path.dirname(__file__), "DoubleColumnCut.pyw")
        if os.path.exists(py_path):
            try:
                subprocess.Popen([sys.executable, py_path])
            except Exception as e:
                messagebox.showerror("Error", f"Failed to open: {str(e)}")
        else:
            messagebox.showerror("Error", "DoubleColumnCut.pyw not found!")

    def load_templates(self):
        """加载模板文件"""
        template_dir = "_Templates"
        try:
            if not os.path.exists(template_dir):
                os.makedirs(template_dir)
                self.log(f"A template catalog has been created: {template_dir}")
                
            templates = [f for f in os.listdir(template_dir) if f.endswith('.docx')]
            if templates:
                self.template_combobox['values'] = templates
                self.template_var.set(templates[0])
                self.config['template_docx'] = os.path.join(template_dir, templates[0])
            else:
                self.log("Tip: There are no template files in the template directory, please place .docx files in the _Templates directory")
        except Exception as e:
            self.log(f"Loading Template Error: {str(e)}")

    def load_macros(self):
        """加载宏文件"""
        macro_dir = "_Macros"
        try:
            if not os.path.exists(macro_dir):
                os.makedirs(macro_dir)
                self.log(f"Macro directory created: {macro_dir}")
                
            macros = [f for f in os.listdir(macro_dir) if f.endswith('.vba')]
            macros.insert(0, "Do not run")  # 默认选项
            self.macro_combobox['values'] = macros
            self.macro_var.set("Do not run")
        except Exception as e:
            self.log(f"Error loading macros: {str(e)}")

    def select_pdf(self):
        """选择PDF文件"""
        file_path = filedialog.askopenfilename(
            filetypes=[("PDF file", "*.pdf"), ("Image files", "*.png;*.jpg;*.jpeg;*.tif;*.tiff;*.bmp;*.gif"), ("All files", "*.*")]
        )
        if file_path:
            self.pdf_path.set(file_path)
            self.log(f"Selected file: {file_path}")

    def select_folder(self):
        """选择包含图片的文件夹"""
        folder = filedialog.askdirectory()
        if folder:
            self.pdf_path.set(folder)
            self.log(f"Selected folder: {folder}")

    def log(self, message, tag=None):
        """记录日志信息"""
        try:
            enabled = bool(getattr(self, 'enable_log_var', tk.BooleanVar(value=False)).get())
        except Exception:
            enabled = False
        if not enabled:
            return

        if threading.current_thread() is not threading.main_thread():
            try:
                self.master.after(0, lambda: self.log(message, tag))
            except Exception:
                pass
            return

        try:
            self.log_area.configure(state="normal")
            self.log_area.insert("end", message + "\n", tag)
            self.log_area.configure(state="disabled")
            self.log_area.see("end")
            self.master.update_idletasks()
        except Exception:
            pass

    def clear_log(self):
        """清空日志"""
        self.log_area.configure(state="normal")
        self.log_area.delete(1.0, "end")
        self.log_area.configure(state="disabled")

    def validate_inputs(self):
        """验证输入参数"""
        try:
            path = self.pdf_path.get()
            if not path:
                raise ValueError("请选择一个 PDF、图片或包含图片的文件夹")
            if not (os.path.isdir(path) or path.lower().endswith(('.pdf', '.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp', '.gif'))):
                raise ValueError("请选择 PDF、图片文件，或包含图片的文件夹")

            template_name = self.template_var.get() if hasattr(self, 'template_var') else ''
            template_path = os.path.join("_Templates", template_name) if template_name else ''
            self.config.update({
                'threshold': int(self.threshold.get()),
                'dpi': int(self.dpi.get()),
                'min_height': int(self.min_height.get()),
                'template_docx': template_path,
                'lr_threshold_enabled': bool(self.lr_enabled_var.get()),
                'lr_threshold_percent': int(self.lr_percent_var.get()),
                'no_insert_word': bool(self.no_insert_var.get()),
                'parallel_enabled': bool(self.parallel_var.get())
            })

            if not (20 <= self.config['threshold'] <= 255):
                raise ValueError("空白阈值需在20-255之间")
            return True
        except Exception as e:
            self.log(f"Input Error: {str(e)}", "error")
            messagebox.showerror("Input Error", str(e))
            return False

    def toggle_advanced(self):
        """显示或隐藏高级选项面板"""
        if getattr(self, 'advanced_visible', False):
            self.advanced_frame.grid_forget()
            self.advanced_visible = False
            self.advanced_btn.config(text="Advanced ▶")
        else:
            self.advanced_frame.grid(row=1, column=0, columnspan=8, pady=(5,0), sticky="w")
            self.advanced_visible = True
            self.advanced_btn.config(text="Advanced ▼")

    def _toggle_log_frame(self):
        """Pack or forget the log frame when user toggles logging in Advanced options."""
        try:
            if getattr(self, 'enable_log_var', tk.BooleanVar(value=False)).get():
                try:
                    self.log_frame.pack(fill="both", expand=True, pady=5, before=self.control_frame)
                except Exception:
                    self.log_frame.pack(fill="both", expand=True, pady=5)
            else:
                try:
                    self.log_frame.pack_forget()
                except Exception:
                    pass
        except Exception:
            pass

    def start_processing(self):
        """启动处理线程"""
        if self.is_processing:
            return

        if not self.validate_inputs():
            return

        self.is_processing = True
        self.process_btn.config(state="disabled")
        self.progress_bar["value"] = 0
        threading.Thread(target=self.process_pdf, daemon=True).start()

    def openf(self):
        main_program_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "__Output")
        try:
            subprocess.Popen(f'explorer "{main_program_dir}"')
        except Exception as e:
            self.log(f"失败啊: {e}")

    def pdf_to_imgs(self, pdf_path, output_dir, render_workers=1, render_log_queue=None):
        """将PDF转换为图片"""
        try:
            img_paths = []

            if os.path.isdir(pdf_path):
                files = sorted(os.listdir(pdf_path))
                image_exts = ('.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp', '.gif')
                imgs = [f for f in files if f.lower().endswith(image_exts)]
                if not imgs:
                    self.log("No image files found in the selected folder", "error")
                    return []
                for i, name in enumerate(imgs):
                    src = os.path.join(pdf_path, name)
                    dst = os.path.join(output_dir, f"page_{i+1}" + os.path.splitext(name)[1])
                    try:
                        shutil.copy2(src, dst)
                    except Exception:
                        try:
                            Image.open(src).save(dst)
                        except Exception as e:
                            self.log(f"Failed to copy/convert image: {src} -> {e}", "error")
                            continue
                    img_paths.append(dst)
                    self.log(f"Collected image {i+1}/{len(imgs)}: {dst}")
                    self.update_progress((i+1)/len(imgs)*30)
                return img_paths

            if pdf_path.lower().endswith(('.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp', '.gif')):
                try:
                    name = os.path.basename(pdf_path)
                    dst = os.path.join(output_dir, name)
                    shutil.copy2(pdf_path, dst)
                    self.log(f"Collected single image: {dst}")
                    self.update_progress(30)
                    return [dst]
                except Exception as e:
                    self.log(f"Failed to copy image file: {e}", "error")
                    return []

            pdf_doc = fitz.open(pdf_path)
            total_pages = len(pdf_doc)

            if total_pages == 0:
                pdf_doc.close()
                return []

            if render_workers and render_workers > 1:
                max_workers = min(render_workers, total_pages)
                pdf_doc.close()
                page_result = {}
                completed = 0
                with concurrent.futures.ProcessPoolExecutor(max_workers=max_workers) as executor:
                    futures = {executor.submit(_render_page_to_image, pdf_path, idx, os.path.join(output_dir, f"page_{idx+1}.png"), int(self.config.get('dpi', 300)), render_log_queue, 0): idx for idx in range(total_pages)}
                    for fut in concurrent.futures.as_completed(futures):
                        try:
                            page_num, out = fut.result()
                            page_result[page_num] = out
                        except Exception as e:
                            try:
                                if render_log_queue is not None:
                                    render_log_queue.put(('error', f"Render worker exception: {e}"))
                            except Exception:
                                pass
                        finally:
                            completed += 1
                            try:
                                self.update_progress((completed / total_pages) * 30)
                            except Exception:
                                pass

                for i in range(1, total_pages + 1):
                    out = page_result.get(i)
                    if out:
                        img_paths.append(out)

                return img_paths
            else:
                for page_num in range(total_pages):
                    page = pdf_doc.load_page(page_num)
                    pix = page.get_pixmap(dpi=self.config['dpi'])
                    output_path = os.path.join(output_dir, f"page_{page_num+1}.png")
                    pix.save(output_path)
                    img_paths.append(output_path)
                    self.log(f"Pages converted {page_num+1}/{total_pages}")
                    self.update_progress((page_num+1)/total_pages*30)

                pdf_doc.close()
                return img_paths
        except Exception as e:
            self.log(f"Error while converting/collecting images: {str(e)}", "error")
            return []

    def process_pdf(self):
        """主处理流程"""
        try:
            pdf_path = self.pdf_path.get()
            base_name = os.path.splitext(os.path.basename(pdf_path))[0]
            output_dir = os.path.join("_temp", f"{base_name}_segments")

            os.makedirs(output_dir, exist_ok=True)
            self.log(f"A temporary directory is created: {output_dir}")

            # 准备日志收集和多进程管理器
            manager = multiprocessing.Manager()
            manager_log_queue = manager.Queue()
            thread_log_queue = queue.Queue()
            self._start_log_collector(thread_log_queue, manager_log_queue)

            render_workers = min(4, os.cpu_count() or 1)
            self.log(f"Start converting/collecting images (process pool, max {render_workers})...")
            old_throttle = getattr(self, '_progress_throttle_seconds', 5)
            self._progress_throttle_seconds = 0.5
            img_paths = self.pdf_to_imgs(pdf_path, output_dir, render_workers=render_workers, render_log_queue=manager_log_queue)
            total_pages = len(img_paths)
            if not img_paths:
                raise ValueError("PDF converting error or no pages/images found")

            cpu = os.cpu_count() or 1
            num_workers = max(1, cpu - 3)
            num_workers = min(num_workers, 8)
            use_parallel = bool(self.config.get('parallel_enabled', False)) and total_pages > 1 and num_workers > 0

            all_segments = []
            segment_counter = 1
            save_interval = int(self.config.get('save_interval', 688))

            # 插入文档的结果路径
            output_docx = None

            if use_parallel:
                self.log(f"Start parallel segmentation of page images using {num_workers} workers...")
                page_items = [(i + 1, img_paths[i]) for i in range(total_pages)]
                num_workers = min(num_workers, total_pages)
                chunk_size = math.ceil(total_pages / num_workers)
                chunks = [page_items[i * chunk_size:(i + 1) * chunk_size] for i in range(num_workers)]
                chunks = [c for c in chunks if c]

                progress_value = manager.Value('i', 0)
                start_time = time.time()

                monitor = threading.Thread(target=self._monitor_progress, args=(total_pages, progress_value, start_time, 30, 50), daemon=True)
                monitor.start()

                with concurrent.futures.ProcessPoolExecutor(max_workers=len(chunks)) as executor:
                    futures = [executor.submit(_worker_segment_images, chunk, self.config, output_dir, int(self.config.get('dpi', 300)), progress_value, idx, manager_log_queue) for idx, chunk in enumerate(chunks)]

                    for fut in concurrent.futures.as_completed(futures):
                        try:
                            result = fut.result()
                            for page_num, seg_paths in result:
                                all_segments.append((page_num, seg_paths))
                        except Exception as e:
                            self.log(f"Parallel worker failed: {e}", "error")

                try:
                    self.master.after(0, lambda: self.update_progress(80))
                except Exception:
                    pass

                all_segments.sort(key=lambda x: x[0])
                sorted_segment_paths = [p for (_, segs) in all_segments for p in segs]

                # 删除原始页面图片
                try:
                    for i in range(len(img_paths)):
                        path = img_paths[i]
                        if path and os.path.exists(path):
                            os.remove(path)
                            self.log(f"Deleted source page image: {path}")
                except Exception:
                    pass

                if not self.config.get('no_insert_word', False):
                    try:
                        self.log("Start inserting all saved segments into the Word document...")
                        output_docx = self._insert_segments_into_doc(sorted_segment_paths, base_name, save_interval=save_interval, base_percent=80)
                    except Exception as e:
                        self.log(f"Insertion failed: {e}", "error")
                else:
                    self.log("Skipped inserting into Word (saved images): all segments kept in temp folder")

            else:
                self.log("Start splitting the picture paragraph (single-threaded)...")
                for idx, img_path in enumerate(img_paths):
                    self.log(f"Processing the picture {idx+1}/{len(img_paths)}: {img_path}")
                    segments, segment_counter = self.save_segments(img_path, output_dir, segment_counter)
                    all_segments.extend(segments)
                    try:
                        os.remove(img_path)
                        self.log(f"Deleted Temporary Files: {img_path}")
                    except Exception:
                        pass
                    self.update_progress(30 + (idx + 1) / len(img_paths) * 50)

                # 循环结束后统一插入
                if not self.config.get('no_insert_word', False):
                    try:
                        self.log("Start inserting all saved segments into the Word document...")
                        output_docx = self._insert_segments_into_doc(all_segments, base_name, save_interval=save_interval, base_percent=80)
                    except Exception as e:
                        self.log(f"Insertion failed: {e}", "error")
                else:
                    self.log("Skipped inserting into Word (saved images): all segments kept in temp folder")

            # 恢复进度节流
            try:
                self._progress_throttle_seconds = old_throttle
            except Exception:
                pass

            # 最终报告
            if not self.config.get('no_insert_word', False):
                if output_docx and os.path.exists(output_docx):
                    final_path = os.path.abspath(output_docx)
                    self.log(f"Processing is complete! The results have been saved to: {final_path}")
                    selected_macro = self.macro_var.get()
                    if selected_macro != "Do not run":
                        macro_path = os.path.join("_Macros", selected_macro)
                        self.inject_and_run_macro(final_path, macro_path)
                    messagebox.showinfo("Processing is complete!", f"The results have been saved to:\n{final_path}")
                else:
                    # 如果由于某种原因没有产出文件，尝试保存模板文档（但正常情况下不应该发生）
                    self.log("Processing complete but no output document found.", "error")
                    messagebox.showerror("Processing Error", "No output document was generated.")
            else:
                self.log("Processing is complete! Images were saved to the temp folder; Word insertion was skipped by configuration.")
                messagebox.showinfo("Processing is complete!", "Images saved; Word insertion skipped as requested.")

        except Exception as e:
            self.log(f"Processing Failure: {str(e)}", "error")
            messagebox.showerror("error", str(e))
        finally:
            self.is_processing = False
            self.master.after(0, lambda: self.process_btn.config(state="normal"))
            self.update_progress(0)

    def inject_and_run_macro(self, docx_path, macro_path):
        """注入并运行选定的宏"""
        try:
            import pythoncom
            pythoncom.CoInitialize()
            with open(macro_path, 'r', encoding='utf-8') as f:
                vba_code = f.read()
            
            macro_name = 'FenScribeMacro'
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            try:
                doc = word.Documents.Open(os.path.abspath(docx_path))
                vba_project = doc.VBProject
                new_module = vba_project.VBComponents.Add(1)
                new_module.CodeModule.AddFromString(vba_code)
                try:
                    word.Application.Run(macro_name)
                except Exception as e:
                    self.log(f"Error running macro: {e}")
                doc.SaveAs2(
                    FileName=os.path.abspath(docx_path),
                    FileFormat=12,
                    AddToRecentFiles=False
                )
            finally:
                doc.Close(SaveChanges=False)
                word.Quit()
            self.log(f"Macro executed and document saved: {docx_path}")
        except Exception as e:
            if -2146822220 in e.args:
                self.log("Please enable VBA project access in Word: File → Options → Trust Center → Trust Center Settings → Macro Settings → Trust access to the VBA project object model")
            else:
                self.log(f"Error injecting macro: {e}")
        finally:
            try:
                pythoncom.CoUninitialize()
            except:
                pass

    def update_progress(self, value):
        """更新进度条"""
        if threading.current_thread() is not threading.main_thread():
            try:
                self.master.after(0, lambda v=value: self.update_progress(v))
            except Exception:
                pass
            return

        now = time.time()
        if value == 0 or value >= 100 or (now - getattr(self, '_last_progress_update', 0)) >= getattr(self, '_progress_throttle_seconds', 5):
            try:
                self.progress_bar["value"] = value
                self.master.update_idletasks()
            except Exception:
                pass
            self._last_progress_update = now
            self._pending_progress = None
            return

        self._pending_progress = value
        if not getattr(self, '_progress_flush_scheduled', False):
            self._progress_flush_scheduled = True
            remaining = int(max(0, (self._progress_throttle_seconds - (now - self._last_progress_update))) * 1000)
            try:
                self.master.after(remaining, self._flush_progress_update)
            except Exception:
                try:
                    self.progress_bar["value"] = value
                    self.master.update_idletasks()
                except Exception:
                    pass

    def save_segments(self, image_path, output_dir, counter):
        """保存分割后的图片片段"""
        try:
            image = Image.open(image_path)
            width, height = image.size
            bottom_margin = int(self.config.get('bottom_margin', 0))

            cropped_image = image.crop((0, 0, width, max(0, height - bottom_margin)))
            segments = _find_segments_from_image(cropped_image, self.config)

            segment_paths = []
            for seg in segments:
                seg_height = seg['end'] - seg['start']
                if seg_height >= int(self.config.get('min_height', 10)):
                    segment_img = cropped_image.crop((0, seg['start'], width, seg['end']))
                    output_path = os.path.join(output_dir, f"segment_{counter}.png")
                    segment_img.save(output_path, dpi=(int(self.config.get('dpi', 300)), int(self.config.get('dpi', 300))))
                    segment_paths.append(output_path)
                    counter += 1

            return segment_paths, counter
        except Exception as e:
            self.log(f"Image Processing Error: {str(e)}", "error")
            return [], counter

    def _insert_segments_into_doc(self, segment_paths, base_name, save_interval=688, base_percent=80, span=20):
        """Insert a list of segment image paths into a single Word document with periodic saves.
        Returns the path to the final saved document.
        """
        try:
            total_segments = len(segment_paths)
            output_dira = os.path.join(os.path.dirname(os.path.abspath(__file__)), "__Output")
            os.makedirs(output_dira, exist_ok=True)
            output_docx = os.path.join(output_dira, f"output_{base_name}.docx")

            try:
                if os.path.exists(output_docx):
                    current_doc = Document(output_docx)
                else:
                    current_doc = Document(self.config['template_docx']) if self.config.get('template_docx') else Document()
            except Exception:
                current_doc = Document(self.config['template_docx']) if self.config.get('template_docx') else Document()

            inserted_count = 0

            for i, seg_path in enumerate(segment_paths):
                try:
                    try:
                        self.master.after(0, lambda idx=i, tot=total_segments: self.progress_label.config(text=f"Inserting to Word: {idx+1}/{tot}"))
                    except Exception:
                        pass

                    self._insert_image_fit_width(current_doc, seg_path)
                    inserted_count += 1

                    if inserted_count % save_interval == 0:
                        try:
                            current_doc.save(output_docx)
                            self.log(f"Saved checkpoint document: {output_docx} (inserts={inserted_count})")
                        except Exception as e:
                            self.log(f"Failed to save checkpoint document: {e}", "error")
                        try:
                            del current_doc
                            gc.collect()
                        except Exception:
                            pass
                        try:
                            current_doc = Document(output_docx)
                        except Exception:
                            current_doc = Document(self.config['template_docx']) if self.config.get('template_docx') else Document()

                except Exception as e:
                    self.log(f"Failed to insert {seg_path}: {e}", "error")

                if total_segments:
                    try:
                        percent = base_percent + (i + 1) / total_segments * span
                        self.update_progress(percent)
                    except Exception:
                        self.update_progress(base_percent)
                else:
                    self.update_progress(100)

            try:
                current_doc.save(output_docx)
                self.log(f"Saved final document: {output_docx}")
                del current_doc
                gc.collect()
            except Exception as e:
                self.log(f"Failed to save final document: {e}", "error")
                return None

            try:
                self.master.after(0, lambda: self.progress_label.config(text=""))
            except Exception:
                pass

            return output_docx
        except Exception as e:
            self.log(f"Error during insertion: {e}", "error")
            return None

    def _flush_progress_update(self):
        """Flush pending progress to the UI (called from main thread)."""
        try:
            self._progress_flush_scheduled = False
            if getattr(self, '_pending_progress', None) is not None:
                val = self._pending_progress
                self._pending_progress = None
                try:
                    self.progress_bar["value"] = val
                    self.master.update_idletasks()
                except Exception:
                    pass
                self._last_progress_update = time.time()
        except Exception:
            pass

    def _monitor_progress(self, total_pages, progress_value, start_time, base=0, span=80):
        """Background monitor that updates the GUI progress and ETA every second."""
        try:
            while getattr(self, 'is_processing', False) and int(progress_value.value) < int(total_pages):
                processed = int(progress_value.value)
                elapsed = time.time() - start_time
                if processed > 0:
                    avg = elapsed / processed
                    eta = int(avg * (total_pages - processed))
                else:
                    eta = None

                frac = (processed / total_pages) if total_pages else 0
                percent = base + min(span, int(frac * span))
                try:
                    self.master.after(0, lambda p=percent, e=eta: self._update_progress_and_label(p, e))
                except Exception:
                    pass
                time.sleep(1)

            try:
                processed = int(progress_value.value)
                frac = (processed / total_pages) if total_pages else 0
                percent = base + min(span, int(frac * span))
                self.master.after(0, lambda p=percent, e=0: self._update_progress_and_label(p, e))
            except Exception:
                pass
        except Exception:
            pass

    def _update_progress_and_label(self, percent, eta_seconds):
        """Update progressbar and ETA label on the main thread."""
        try:
            self.update_progress(percent)
            if eta_seconds is None:
                eta_text = "Estimating..."
            else:
                eta_text = str(timedelta(seconds=int(eta_seconds)))
            self.progress_label.config(text=f"{percent:.0f}% — ETA: {eta_text}")
        except Exception:
            pass

    def _start_log_collector(self, thread_queue=None, manager_queue=None):
        """Start a background thread that polls provided queues and forwards messages to the GUI log."""
        def _collector():
            while getattr(self, 'is_processing', False):
                if thread_queue is not None:
                    try:
                        item = thread_queue.get(timeout=0.5)
                        if item:
                            level, msg = item
                            if level == 'info':
                                try:
                                    self.master.after(0, lambda m=msg: self.log(m))
                                except Exception:
                                    pass
                            else:
                                try:
                                    self.master.after(0, lambda m=msg: self.log(m, 'error'))
                                except Exception:
                                    pass
                    except Exception:
                        pass

                if manager_queue is not None:
                    try:
                        item = manager_queue.get(timeout=0.5)
                        if item:
                            try:
                                level, msg = item
                                if level == 'info':
                                    self.master.after(0, lambda m=msg: self.log(m))
                                else:
                                    self.master.after(0, lambda m=msg: self.log(m, 'error'))
                            except Exception:
                                pass
                    except Exception:
                        pass

            for q in (thread_queue, manager_queue):
                if q is None:
                    continue
                while True:
                    try:
                        item = q.get_nowait() if hasattr(q, 'get_nowait') else q.get(timeout=0.1)
                        if item:
                            try:
                                level, msg = item
                                if level == 'info':
                                    self.master.after(0, lambda m=msg: self.log(m))
                                else:
                                    self.master.after(0, lambda m=msg: self.log(m, 'error'))
                            except Exception:
                                pass
                    except Exception:
                        break

        threading.Thread(target=_collector, daemon=True).start()

    def _insert_image_fit_width(self, doc, img_path):
        """Insert an image into `doc` scaled to the document's page width (margins considered)."""
        section = doc.sections[0]
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin
        available_width = page_width - left_margin - right_margin

        try:
            sectPr = section._sectPr
            cols_elem = sectPr.xpath('./w:cols')
            col_count = None
            if cols_elem:
                num_attr = cols_elem[0].get(qn('w:num'))
                if num_attr:
                    try:
                        col_count = int(num_attr)
                    except Exception:
                        col_count = None
                if col_count is None:
                    col_elems = cols_elem[0].xpath('./w:col')
                    if col_elems:
                        col_count = len(col_elems)
                if col_count and col_count > 1:
                    available_width = int(available_width / col_count)
        except Exception:
            pass

        with Image.open(img_path) as img:
            img_width_px, img_height_px = img.size
            dpi = self.config.get('dpi', 300)
            if dpi <= 0:
                dpi = 300
            img_width_in = img_width_px / dpi
            img_height_in = img_height_px / dpi

        emu_per_inch = 914400
        avail_in_inch = available_width / emu_per_inch
        target_width_in = min(img_width_in, avail_in_inch)

        doc.add_picture(img_path, width=Inches(target_width_in))

if __name__ == "__main__":
    root = TkinterDnD.Tk()  
    app = PDFProcessorApp(root)
    try:
        root.iconbitmap('icon.ico')
        ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID("com.ordylan.fenscribe")
    except Exception as e:
        print("Error loading icon:", e)

    if len(sys.argv) > 1:
        args = sys.argv[1:]
        image_exts = ('.pdf', '.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp', '.gif')
        valid = [p for p in args if os.path.isdir(p) or p.lower().endswith(image_exts)]
        if not valid:
            messagebox.showerror("Error", "Please pass a PDF, image, or folder containing images as a parameter")
        else:
            chosen = valid[0]
            if len(valid) > 1:
                app.log(f"Multiple inputs passed; using first: {chosen}")
            app.pdf_path.set(chosen)
            app.log(f"File was loaded via the command line: {chosen}")

    style = ttk.Style()
    style.theme_use("clam")
    style.configure("TButton", padding=6)
    style.configure("TEntry", padding=5)
    style.configure("TCombobox", padding=5)
    os.makedirs("__Output", exist_ok=True)
    root.mainloop()